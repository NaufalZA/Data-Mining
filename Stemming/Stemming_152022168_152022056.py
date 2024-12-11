import csv
import re
from tkinter import filedialog
import tkinter as tk
import PyPDF2
from docx import Document
import os

class Stemmer:
    def __init__(self):
        self.kamus = self.load_kamus()
        self.stopwords = self.load_stopwords()
        self.forbidden_combinations = {
            'be': ['i'],
            'di': ['an'],
            'ke': ['i', 'kan'],
            'me': ['an'],
            'se': ['i', 'kan']
        }

    def load_kamus(self):
        with open('documents/Kamus.txt', 'r') as file:
            return set(word.strip().lower() for word in file)

    def load_stopwords(self):
        stopwords = set()
        with open('documents/Stopword.csv', 'r') as file:
            reader = csv.reader(file)
            for row in reader:
                stopwords.add(row[0].lower())
        return stopwords

    def is_valid_word(self, word):
        return bool(re.match('^[a-zA-Z]+$', word))

    def remove_stopwords(self, text):
        words = text.lower().split()
        return ' '.join([word for word in words if word not in self.stopwords])

    def check_kamus(self, word):
        return word.lower() in self.kamus

    def remove_inflection_suffixes(self, word):
        # Remove -lah, -kah, -tah, -pun
        if word.endswith(('lah', 'kah', 'tah', 'pun')):
            word = re.sub(r'(lah|kah|tah|pun)$', '', word)
            
        # Remove -ku, -mu, -nya
        if word.endswith(('ku', 'mu', 'nya')):
            word = re.sub(r'(ku|mu|nya)$', '', word)
            
        return word

    def remove_derivation_suffixes(self, word):
        if word.endswith(('i', 'an', 'kan')):
            original = word
            
            # Remove -i, -an, or -kan
            if word.endswith('kan'):
                word = word[:-3]
            elif word.endswith(('i', 'an')):
                word = word[:-2]

            # Special case for -an where last letter is k
            if original.endswith('an') and word.endswith('k'):
                word = word[:-1]
                if self.check_kamus(word):
                    return word
                word = word + 'k'

            # Check if result exists in dictionary
            if self.check_kamus(word):
                return word
                
            # If not found, restore the original suffix
            return original
            
        return word

    def is_forbidden_combination(self, prefix, suffix):
        forbidden = {
            'be': ['i'],
            'di': ['an'],
            'ke': ['i', 'kan'],
            'me': ['an'],
            'se': ['i', 'kan']
        }
        return prefix in forbidden and suffix in forbidden[prefix]

    def remove_prefix(self, word, iteration=1):
        if iteration > 3:
            return word

        if word[:2] in ['di', 'ke', 'se']:
            prefix = word[:2]
            stemmed = word[2:]
            
        elif word.startswith(('ter', 'bel')):
            prefix = word[:3]
            stemmed = word[3:]
            
        elif word.startswith(('me', 'pe', 'be')):
            if len(word) > 3:
                if word[2] == 'r' and word[3] in ['a','i','u','e','o']:
                    prefix = word[:3]
                    stemmed = word[3:]
                else:
                    prefix = word[:2]
                    stemmed = word[2:]
            else:
                prefix = word[:2]
                stemmed = word[2:]
        else:
            return word

        # Check if result exists in dictionary
        if self.check_kamus(stemmed):
            return stemmed
            
        # Try next iteration
        return self.remove_prefix(word, iteration + 1)

    def stem_word(self, word):
        steps = []
        if not self.is_valid_word(word):
            return word, steps
            
        # Step 1: Check in dictionary
        steps.append(f"Step 1: Checking '{word}' in dictionary")
        if self.check_kamus(word):
            steps.append("Result: Found in dictionary")
            return word, steps

        original_word = word
        suffix_removed = False

        # Step 2: Remove inflection suffixes
        steps.append(f"Step 2: Checking inflection suffixes for '{word}'")
        if any(word.endswith(suffix) for suffix in ['lah', 'kah', 'tah', 'pun']):
            old_word = word
            word = self.remove_inflection_suffixes(word)
            steps.append(f"Removed particle suffix: '{old_word}' → '{word}'")
            
            if any(word.endswith(suffix) for suffix in ['ku', 'mu', 'nya']):
                old_word = word
                word = self.remove_inflection_suffixes(word)
                steps.append(f"Removed possessive pronoun: '{old_word}' → '{word}'")

        if self.check_kamus(word):
            steps.append(f"Result: Found '{word}' in dictionary after inflection removal")
            return word, steps

        # Step 3: Remove derivation suffixes
        steps.append(f"Step 3: Checking derivation suffixes for '{word}'")
        if any(word.endswith(suffix) for suffix in ['i', 'an', 'kan']):
            temp_word = word
            if word.endswith('kan'):
                word = word[:-3]
                suffix_removed = 'kan'
                steps.append(f"Removed -kan: '{temp_word}' → '{word}'")
            elif word.endswith('i'):
                word = word[:-1]
                suffix_removed = 'i'
                steps.append(f"Removed -i: '{temp_word}' → '{word}'")
            elif word.endswith('an'):
                word = word[:-2]
                suffix_removed = 'an'
                steps.append(f"Removed -an: '{temp_word}' → '{word}'")
                
                if word.endswith('k'):
                    k_word = word[:-1]
                    steps.append(f"Checking k-removal: '{word}' → '{k_word}'")
                    if self.check_kamus(k_word):
                        steps.append(f"Result: Found '{k_word}' in dictionary")
                        return k_word, steps
                    word = temp_word
                    steps.append("Restored original: k-removal unsuccessful")

            if self.check_kamus(word):
                steps.append(f"Result: Found '{word}' in dictionary")
                return word, steps
            word = temp_word
            steps.append("Restored original: suffix removal unsuccessful")

        # Step 4: Check prefix-suffix combination
        if suffix_removed:
            prefix = self.get_prefix_type(word)
            steps.append(f"Step 4: Checking prefix-suffix combination: prefix='{prefix}', suffix='{suffix_removed}'")
            if prefix and suffix_removed in self.forbidden_combinations.get(prefix, []):
                steps.append(f"Found forbidden combination: {prefix}- with -{suffix_removed}")
                return original_word, steps

        # Step 4b: Remove prefixes
        steps.append(f"Step 4b: Removing prefixes from '{word}'")
        result = self.remove_prefix(word)
        if result != word:
            steps.append(f"Removed prefix: '{word}' → '{result}'")
            if self.check_kamus(result):
                steps.append(f"Result: Found '{result}' in dictionary")
                return result, steps
        
        # Step 6: Return original word if no root found
        steps.append("Step 6: No root word found, returning original word")
        return original_word, steps

    def get_prefix_type(self, word):
        if word.startswith(('di', 'ke', 'se')):
            return word[:2]
        elif word.startswith(('ter', 'bel')):
            return word[:3]
        elif word.startswith(('me', 'pe', 'be')):
            return word[:2]
        return None

    def stem_text(self, text):
        # Remove stopwords first
        text = self.remove_stopwords(text)
        
        # Split into words and stem each word
        words = text.split()
        results = []
        all_steps = []
        
        for word in words:
            stemmed_word, steps = self.stem_word(word)
            results.append(stemmed_word)
            all_steps.append((word, stemmed_word, steps))
            
        return ' '.join(results), all_steps

def read_file_content(file_path):
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ''
            for page in pdf_reader.pages:
                text += page.extract_text() + '\n'
            return text
            
    elif file_extension == '.docx':
        doc = Document(file_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        return text
        
    elif file_extension == '.txt':
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    else:
        raise ValueError("Unsupported file format")

def export_to_word(original_text, stemmed_text, steps, input_file_path):
    original_filename = os.path.splitext(os.path.basename(input_file_path))[0]
    output_filename = f"results/Stemmed_{original_filename}.docx"
    
    doc = Document()
    doc.add_heading('Stemming Results', 0)
    
    # Add stemming process details
    doc.add_heading('Stemming Process:', level=1)
    for original, stemmed, word_steps in steps:
        if original != stemmed:  # Only show words that were actually stemmed
            doc.add_heading(f"Word: {original} → {stemmed}", level=2)
            for step in word_steps:
                doc.add_paragraph(step, style='List Bullet')
            doc.add_paragraph()  # Add space between words
    
    # Add final stemmed text
    doc.add_heading('Final Stemmed Text:', level=1)
    doc.add_paragraph(stemmed_text)
    
    doc.save(output_filename)
    return output_filename

if __name__ == "__main__":
    stemmer = Stemmer()
    
    root = tk.Tk()
    root.withdraw()
    
    file_path = filedialog.askopenfilename(
        title="Select File",
        filetypes=[
            ("All supported files", "*.txt;*.pdf;*.docx"),
            ("Text files", "*.txt"),
            ("PDF files", "*.pdf"),
            ("Word files", "*.docx"),
            ("All files", "*.*")
        ]
    )
    
    if file_path:
        try:
            text = read_file_content(file_path)
            stemmed_text, steps = stemmer.stem_text(text)
            
            print(f"\nStemmed text:")
            print(stemmed_text)
            
            # Export results to Word document
            output_file = export_to_word(text, stemmed_text, steps, file_path)
            print(f"\nResults exported to: {output_file}")
            
        except Exception as e:
            print(f"Error processing file: {e}")
    else:
        print("No file selected")
