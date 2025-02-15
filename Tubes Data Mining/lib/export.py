from docx import Document
import os
import math
from stem import Stemmer

def export_to_word(original_text, stemmed_text, steps, input_file_path, vsm_data=None, stemmer=None, file_paths=None):
    if stemmer is None:
        stemmer = Stemmer()
        
    original_filename = os.path.splitext(os.path.basename(input_file_path))[0]
    output_filename = f"results/Hasil_{original_filename}.docx"
    
    doc = Document()
    doc.add_heading('Hasil Stemming dan Perhitungan Temu Balik', 0)
    
    doc.add_heading('Original Text:', level=1)
    doc.add_paragraph(original_text)
    
    tokens = stemmer.tokenize(original_text)
    doc.add_heading('1. Hasil Tokenisasi:', level=1)
    token_text = ', '.join([t['token'] for t in tokens])
    doc.add_paragraph(token_text)

    doc.add_heading('2. Hasil Setelah Removal Stopword:', level=1)
    tokens_without_stopwords = [t for t in tokens if t['token'] not in stemmer.stopwords]
    stopword_text = ', '.join([t['token'] for t in tokens_without_stopwords])
    doc.add_paragraph(stopword_text)
    
    doc.add_heading('3. Hasil Pengecekan Kata Dalam Kamus:', level=1)
    dict_check = [(t['token'], stemmer.check_kamus(t['token'])) for t in tokens_without_stopwords]
    
    total_words = len(dict_check)
    found_words = sum(1 for _, in_dict in dict_check if in_dict)
    
    summary = doc.add_paragraph()
    summary.add_run(f"Total kata: {total_words}\n").bold = True
    summary.add_run(f"Kata dasar dalam kamus: {found_words}\n").bold = True
    
    dict_table = doc.add_table(rows=1, cols=2)
    dict_table.style = 'Table Grid'
    header_cells = dict_table.rows[0].cells
    header_cells[0].text = 'Kata'
    header_cells[1].text = 'Status Kamus'
    
    for word, in_dict in dict_check:
        row_cells = dict_table.add_row().cells
        row_cells[0].text = word
        row_cells[1].text = 'Ada' if in_dict else 'Tidak ada'

    doc.add_heading('4. Proses Stemming Step by Step:', level=1)
    if steps:
        successful_stems = []
        failed_stems = []
        for step in steps:
            original, stemmed, word_steps = step
            if stemmed != original: 
                successful_stems.append(step)
            else: 
                failed_stems.append(step)
        
        steps_to_show = successful_stems[:8] + failed_stems[:2]
        
        for original, stemmed, word_steps in steps_to_show:
            p = doc.add_paragraph()
            p.add_run(f"Kata: {original}\n").bold = True
            
            # Langkah 1: Cek Kata Dasar Awal
            p.add_run("1. Pengecekan Kata Dasar Awal: ")
            if stemmer.check_kamus(original):
                p.add_run("Ditemukan dalam kamus → Proses stemming berhenti\n")
                p.add_run(f"Hasil akhir: {original}\n")
                p.add_run("─" * 40 + "\n")
                continue
            else:
                p.add_run("Tidak ditemukan dalam kamus → Lanjut proses stemming\n")
            
            current_word = original
            p.add_run("2. Removal Inflectional Suffix: ")
            inflection_removed, inflection_suffix = stemmer.remove_inflection_suffixes(current_word)
            if inflection_removed != current_word:
                p.add_run(f"{current_word} → {inflection_removed}")
                if inflection_suffix:
                    p.add_run(f" (Suffix: {inflection_suffix})\n")
                else:
                    p.add_run("\n")
                if stemmer.check_kamus(inflection_removed):
                    p.add_run("Hasil removal Inflectional Suffix ditemukan dalam kamus → Proses stemming berhenti\n")
                    p.add_run(f"Hasil akhir: {inflection_removed}\n")
                    p.add_run("─" * 40 + "\n")
                    continue
                p.add_run("Hasil removal Inflectional Suffix tidak ditemukan dalam kamus → Lanjut proses stemming\n")
                current_word = inflection_removed
            else:
                p.add_run("Tidak ada Inflectional Suffix\n")
            
            p.add_run("3. Removal Derivational Suffix: ")
            derivation_removed, derivation_suffix = stemmer.remove_derivation_suffixes(current_word)
            if derivation_removed != current_word:
                p.add_run(f"{current_word} → {derivation_removed}")
                if derivation_suffix:
                    p.add_run(f" (Suffix: {derivation_suffix})\n")
                else:
                    p.add_run("\n")
                if stemmer.check_kamus(derivation_removed):
                    p.add_run("Hasil removal Derivational Suffix ditemukan dalam kamus → Proses stemming berhenti\n")
                    p.add_run(f"Hasil akhir: {derivation_removed}\n")
                    p.add_run("─" * 40 + "\n")
                    continue
                p.add_run("Hasil removal Derivational Suffix tidak ditemukan dalam kamus → Lanjut proses stemming\n")
                current_word = derivation_removed
            else:
                p.add_run("Tidak ada Derivational Suffix\n")
            
            p.add_run("4. Removal Prefix: ")
            prefix_removed, prefix_type = stemmer.remove_prefix(current_word)
            if prefix_type:
                p.add_run(f"{current_word} → {prefix_removed} (Prefix: {prefix_type})\n")
                if stemmer.check_kamus(prefix_removed):
                    p.add_run("Hasil removal Prefix ditemukan dalam kamus → Proses stemming berhenti\n")
                else:
                    p.add_run("Hasil removal Prefix tidak ditemukan dalam kamus\n")
            else:
                p.add_run("Tidak ada Prefix\n")
            
            p.add_run(f"\nHasil akhir: {stemmed}\n")
            p.add_run("─" * 40 + "\n")

    doc.add_heading('5. Hasil Akhir Stemming:', level=1)
    doc.add_paragraph(stemmed_text)
    
    if vsm_data:
        doc.add_heading('6. Analisis VSM:', level=1)
        
        doc.add_paragraph("Berikut adalah matriks term-dokumen yang menunjukkan frekuensi kemunculan setiap kata pada setiap dokumen:")
        
        terms = sorted(list(vsm_data['terms']))
        table = doc.add_table(rows=1, cols=len(vsm_data['documents']) + 1)
        table.style = 'Table Grid'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Term'
        for i in range(len(vsm_data['documents'])):
            header_cells[i+1].text = f'Doc {i+1}'
            
        for term in terms:
            row_cells = table.add_row().cells
            row_cells[0].text = term
            for doc_id in range(len(vsm_data['documents'])):
                freq = vsm_data['term_doc_freq'][term][doc_id]
                row_cells[doc_id+1].text = str(freq)

        doc.add_heading('Inverse Document Frequency (IDF):', level=2)
        idf_table = doc.add_table(rows=1, cols=2)
        idf_table.style = 'Table Grid'
        idf_table.rows[0].cells[0].text = 'Term'
        idf_table.rows[0].cells[1].text = 'IDF'
        
        for term in sorted(vsm_data['terms']):
            row_cells = idf_table.add_row().cells
            row_cells[0].text = term
            row_cells[1].text = f"{vsm_data['idf'][term]:.2f}".replace('.', ',')

        doc.add_heading('TF-IDF Weights:', level=2)
        tf_idf_table = doc.add_table(rows=1, cols=len(vsm_data['documents']) + 1)
        tf_idf_table.style = 'Table Grid'
        
        header_cells = tf_idf_table.rows[0].cells
        header_cells[0].text = 'Term'
        for i in range(len(vsm_data['documents'])):
            header_cells[i+1].text = f'Doc {i+1}'
            
        for term in sorted(list(vsm_data['terms'])):
            row_cells = tf_idf_table.add_row().cells
            row_cells[0].text = term
            for doc_id in range(len(vsm_data['documents'])):
                tf_idf = vsm_data['tf_idf_vectors'][doc_id].get(term, 0)
                row_cells[doc_id+1].text = f"{tf_idf:.2f}".replace('.', ',')

    if vsm_data and 'query' in vsm_data:
        doc.add_heading('7. Hasil Pencarian dan Perhitungan Similiarity:', level=1)
        doc.add_paragraph(f"Kata Kunci: {vsm_data['query']}")
        
        doc.add_heading('Analisis Keyword:', level=2)
        doc.add_paragraph("Frekuensi kemunculan kata pada query setelah preprocessing:")
        
        stemmed_query, _ = stemmer.stem_text(vsm_data['query'])
        query_terms = stemmed_query.split()
        query_vector = {}
        for term in query_terms:
            if term in vsm_data['terms']:
                query_vector[term] = query_terms.count(term)
                
        query_table = doc.add_table(rows=1, cols=2)
        query_table.style = 'Table Grid'
        query_table.rows[0].cells[0].text = 'Term'
        query_table.rows[0].cells[1].text = 'Frequency'
        for term, freq in query_vector.items():
            row = query_table.add_row()
            row.cells[0].text = term
            row.cells[1].text = str(freq)
            
        doc.add_heading('Perhitungan Similiarity:', level=2)
        for i, (result_doc, score) in enumerate(vsm_data['search_results']):
            if score > 0:
                doc_id = result_doc['id'] + 1
                doc.add_paragraph(f"\nSimilarity (D{doc_id}, Q) = cos(D{doc_id}, Q)")
                
                doc_vector = vsm_data['doc_vectors'][result_doc['id']]
                query_vector = {term: query_terms.count(term) for term in query_terms if term in vsm_data['terms']}
                
                doc.add_paragraph("Perhitungan Dot:")
                dot_products = []
                dot_product = 0
                for term in vsm_data['terms']:
                    q_val = query_vector.get(term, 0)
                    d_val = doc_vector.get(term, 0)
                    if q_val > 0 or d_val > 0:
                        dot_products.append(f"({q_val} × {d_val})")
                        dot_product += q_val * d_val
                
                doc.add_paragraph("Dot(D{}, Q) = {}".format(
                    doc_id,
                    " + ".join(dot_products) + f" = {dot_product}"
                ))
                
                doc.add_paragraph("\nPerhitungan Magnitude :")
                q_magnitude_parts = []
                for term, val in query_vector.items():
                    if val > 0:
                        q_magnitude_parts.append(f"{val}²")
                query_magnitude = math.sqrt(sum(v * v for v in query_vector.values()))
                doc.add_paragraph("|Q| = √({}) = {:.2f}".format(
                    " + ".join(q_magnitude_parts),
                    query_magnitude
                ).replace('.', ','))
                
                d_magnitude_parts = []
                for term, val in doc_vector.items():
                    if val > 0:
                        d_magnitude_parts.append(f"{val}²")
                doc_magnitude = math.sqrt(sum(v * v for v in doc_vector.values()))
                doc.add_paragraph("|D{}| = √({}) = {:.2f}".format(
                    doc_id,
                    " + ".join(d_magnitude_parts),
                    doc_magnitude
                ).replace('.', ','))
                
                doc.add_paragraph("\nSimilarity(D{}, Q) = Dot(D{}, Q) / (|D{}| × |Q|)".format(
                    doc_id, doc_id, doc_id
                ))
                doc.add_paragraph("Similarity(D{}, Q) = {} / ({:.2f} × {:.2f})".format(
                    doc_id, dot_product, doc_magnitude, query_magnitude
                ).replace('.', ','))
                
                doc.add_paragraph(f"Similarity(D{doc_id}, Q) = {score:.2f}".replace('.', ','))
                
                doc.add_paragraph("-" * 40)

        doc.add_heading('Ranking Documents:', level=2)
        results_table = doc.add_table(rows=1, cols=2)
        results_table.style = 'Table Grid'
        header_cells = results_table.rows[0].cells
        header_cells[0].text = 'Judul Dokumen'
        header_cells[1].text = 'Nilai Similiarity'
        
        for result_doc, score in vsm_data['search_results']:
            if score > 0 and file_paths:
                row_cells = results_table.add_row().cells
                row_cells[0].text = os.path.basename(file_paths[result_doc['id']])
                row_cells[1].text = f"{score:.2f}".replace('.', ',')

    doc.save(output_filename)
    return output_filename