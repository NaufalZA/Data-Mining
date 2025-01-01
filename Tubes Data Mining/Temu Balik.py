import csv
import re
from tkinter import filedialog
import tkinter as tk
import PyPDF2
from docx import Document
import os
import string
import math
from collections import defaultdict

class Stemmer:
    def __init__(self):
        self.kamus = self.load_kamus()
        self.stopwords = self.load_stopwords()
        self.forbidden_combinations = {
            'be': ['i'],
            'di': ['an'],
            'ke': ['i', 'kan'],
            'me': ['an'],
            'se': ['i', 'kan']
        }
        self.punctuation = string.punctuation + '"' + '"' + ''' + ''' + '—' + '–'
        self.prefix_types = {
            'di': 'di-',
            'ke': 'ke-',
            'se': 'se-',
            'te': 'te-',
            'ter': 'ter-',
            'me': 'me-',
            'be': 'be-',
            'pe': 'pe-'
        }

    def load_kamus(self):
        with open('documents/Kamus.txt', 'r') as file:
            return set(word.strip().lower() for word in file)

    def load_stopwords(self):
        stopwords = set()
        with open('documents/Stopword.csv', 'r') as file:
            reader = csv.reader(file)
            for row in reader:
                stopwords.add(row[0].lower())
        return stopwords

    def is_valid_word(self, word):
        return bool(re.match('^[a-zA-Z]+$', word))

    def remove_stopwords(self, text):
        words = text.lower().split()
        return ' '.join([word for word in words if word not in self.stopwords])

    def check_kamus(self, word):
        return word.lower() in self.kamus

    def remove_inflection_suffixes(self, word):
        # Remove -lah, -kah, -tah, -pun
        if word.endswith(('lah', 'kah', 'tah', 'pun')):
            word = re.sub(r'(lah|kah|tah|pun)$', '', word)
            
        # Remove -ku, -mu, -nya
        if word.endswith(('ku', 'mu', 'nya')):
            word = re.sub(r'(ku|mu|nya)$', '', word)
            
        return word

    def remove_derivation_suffixes(self, word):
        if word.endswith(('i', 'an', 'kan')):
            original = word
            
            # Remove -i, -an, or -kan
            if word.endswith('kan'):
                word = word[:-3]
            elif word.endswith(('i', 'an')):
                word = word[:-2]

            # Special case for -an where last letter is k
            if original.endswith('an') and word.endswith('k'):
                word = word[:-1]
                if self.check_kamus(word):
                    return word
                word = word + 'k'

            # Check if result exists in dictionary
            if self.check_kamus(word):
                return word
                
            # If not found, restore the original suffix
            return original
            
        return word

    def is_forbidden_combination(self, prefix, suffix):
        forbidden = {
            'be': ['i'],
            'di': ['an'],
            'ke': ['i', 'kan'],
            'me': ['an'],
            'se': ['i', 'kan']
        }
        return prefix in forbidden and suffix in forbidden[prefix]

    def remove_prefix(self, word, iteration=1):
        if iteration > 3:
            return word, None  # Return None as prefix type if max iterations reached

        # Track previous prefix for comparisons
        previous_prefix = None if iteration == 1 else self.get_prefix_type(word)

        if word[:2] in ['di', 'ke', 'se']:
            prefix = word[:2]
            stemmed = word[2:]
            prefix_type = self.prefix_types[prefix]
            
        elif word.startswith('ter'):
            prefix = 'ter'
            stemmed = word[3:]
            prefix_type = 'ter-'
            
        elif word.startswith('te'):
            prefix = 'te'
            stemmed = word[2:]
            prefix_type = 'te-'
            
        elif word.startswith(('me', 'pe', 'be')):
            if len(word) > 3 and word[2] == 'r' and word[3] in 'aiueo':
                prefix = word[:3]
                stemmed = word[3:]
            else:
                prefix = word[:2]
                stemmed = word[2:]
            prefix_type = self.prefix_types[prefix[:2]]
        else:
            return word, None

        # Stop if same prefix is found twice
        if previous_prefix == prefix_type:
            return word, None

        # Check dictionary
        if self.check_kamus(stemmed):
            return stemmed, prefix_type

        # Try recoding for certain prefixes
        recoded = self.recode_prefix(prefix, stemmed)
        if recoded != stemmed and self.check_kamus(recoded):
            return recoded, prefix_type
            
        # Try next iteration
        next_word, next_prefix = self.remove_prefix(word, iteration + 1)
        return next_word, next_prefix or prefix_type

    def recode_prefix(self, prefix, word):
        """Handle special recoding cases"""
        if prefix in ['me', 'pe']:
            if word.startswith('ng'):
                return word[2:]  # ng -> ''
            elif word.startswith('ny'):
                return 's' + word[2:]  # ny -> s
            elif word.startswith('n'):
                return 't' + word[1:]  # n -> t
        return word

    def stem_word(self, word):
        steps = []
        if not self.is_valid_word(word):
            return word, steps
            
        # Step 1: Check in dictionary
        steps.append(f"Step 1: Checking '{word}' in dictionary")
        if self.check_kamus(word):
            steps.append("Result: Found in dictionary")
            return word, steps

        original_word = word
        suffix_removed = False
        prefix_type = None

        # Step 2: Remove inflection suffixes with detailed tracking
        steps.append("Step 2: Checking inflection suffixes")
        if any(word.endswith(suffix) for suffix in ['lah', 'kah', 'tah', 'pun', 'ku', 'mu', 'nya']):
            old_word = word
            word = self.remove_inflection_suffixes(word)
            steps.append(f"Removed inflection suffix: {old_word} → {word}")

        # Step 3: Remove derivation suffixes
        steps.append(f"Step 3: Checking derivation suffixes for '{word}'")
        if any(word.endswith(suffix) for suffix in ['i', 'an', 'kan']):
            temp_word = word
            if word.endswith('kan'):
                word = word[:-3]
                suffix_removed = 'kan'
                steps.append(f"Removed -kan: '{temp_word}' → '{word}'")
            elif word.endswith('i'):
                word = word[:-1]
                suffix_removed = 'i'
                steps.append(f"Removed -i: '{temp_word}' → '{word}'")
            elif word.endswith('an'):
                word = word[:-2]
                suffix_removed = 'an'
                steps.append(f"Removed -an: '{temp_word}' → '{word}'")
                
                if word.endswith('k'):
                    k_word = word[:-1]
                    steps.append(f"Checking k-removal: '{word}' → '{k_word}'")
                    if self.check_kamus(k_word):
                        steps.append(f"Result: Found '{k_word}' in dictionary")
                        return k_word, steps
                    word = temp_word
                    steps.append("Restored original: k-removal unsuccessful")

            if self.check_kamus(word):
                steps.append(f"Result: Found '{word}' in dictionary")
                return word, steps
            word = temp_word
            steps.append("Restored original: suffix removal unsuccessful")

        # Step 4: Prefix removal with better tracking
        if suffix_removed:
            steps.append("Step 4a: Checking prefix-suffix combinations")
            prefix = self.get_prefix_type(word)
            if prefix and suffix_removed in self.forbidden_combinations.get(prefix, []):
                steps.append(f"Found forbidden combination: {prefix}- with -{suffix_removed}")
                return original_word, steps
        
        steps.append("Step 4b: Removing prefixes")
        word, prefix_type = self.remove_prefix(word)
        if prefix_type:
            steps.append(f"Removed prefix type: {prefix_type}")

        # Step 5: Recoding (if needed)
        if prefix_type:
            steps.append("Step 5: Checking recoding rules")
            recoded = self.recode_prefix(prefix_type.rstrip('-'), word)
            if recoded != word:
                steps.append(f"Applied recoding: {word} → {recoded}")
                word = recoded

        # Step 6: Return original word if no root found
        steps.append("Step 6: No root word found, returning original word")
        return original_word, steps

    def get_prefix_type(self, word):
        if word.startswith(('di', 'ke', 'se')):
            return word[:2]
        elif word.startswith(('ter', 'bel')):
            return word[:3]
        elif word.startswith(('me', 'pe', 'be')):
            return word[:2]
        return None

    def tokenize(self, text):
        """
        Tokenize text into words while handling punctuation and special cases.
        Returns list of tokens and their positions.
        """
        # Convert to lowercase
        text = text.lower()
        
        # Replace newlines with spaces
        text = text.replace('\n', ' ')
        
        # Handle punctuation
        for p in self.punctuation:
            text = text.replace(p, ' ')
            
        # Remove extra whitespace
        text = ' '.join(text.split())
        
        # Create tokens with positions
        words = text.split()
        tokens = []
        position = 0
        
        for word in words:
            # Skip empty strings and purely numeric tokens
            if word and not word.isnumeric():
                tokens.append({
                    'token': word,
                    'position': position,
                    'original': word
                })
            position += 1
            
        return tokens

    def stem_text(self, text):
        # First tokenize the text
        tokens = self.tokenize(text)
        
        # Remove stopwords
        tokens = [t for t in tokens if t['token'] not in self.stopwords]
        
        results = []
        all_steps = []
        
        # Process each token
        for token in tokens:
            stemmed_word, steps = self.stem_word(token['token'])
            token['stemmed'] = stemmed_word
            results.append(stemmed_word)
            if token['token'] != stemmed_word:  # Only track changes
                all_steps.append((token['original'], stemmed_word, steps))
            
        return ' '.join(results), all_steps

def read_file_content(file_path):
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ''
            for page in pdf_reader.pages:
                text += page.extract_text() + '\n'
            return text
            
    elif file_extension == '.docx':
        doc = Document(file_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        return text
        
    elif file_extension == '.txt':
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    else:
        raise ValueError("Unsupported file format")

def export_to_word(original_text, stemmed_text, steps, input_file_path, vsm_data=None):
    original_filename = os.path.splitext(os.path.basename(input_file_path))[0]
    output_filename = f"results/Stemmed_{original_filename}.docx"
    
    doc = Document()
    doc.add_heading('Stemming and VSM Results', 0)
    
    # Show tokenization results
    tokens = stemmer.tokenize(original_text)
    doc.add_heading('1. After Tokenization:', level=1)
    token_text = ', '.join([t['token'] for t in tokens])
    doc.add_paragraph(token_text)
    
    # Show after stopword removal
    filtered_tokens = [t for t in tokens if t['token'] not in stemmer.stopwords]
    doc.add_heading('2. After Stopword Removal:', level=1)
    stopword_text = ', '.join([t['token'] for t in filtered_tokens])
    doc.add_paragraph(stopword_text)
    
    # Show after number removal
    valid_tokens = [t for t in filtered_tokens if not any(c.isdigit() for c in t['token'])]
    doc.add_heading('3. After Number Removal:', level=1)
    valid_text = ', '.join([t['token'] for t in valid_tokens])
    doc.add_paragraph(valid_text)
    
    # Show words found in dictionary with counts
    dict_check = [(t['token'], stemmer.check_kamus(t['token'])) for t in valid_tokens]
    doc.add_heading('4. Dictionary Check:', level=1)
    
    # Count total words and found root words
    total_words = len(dict_check)
    found_words = sum(1 for _, in_dict in dict_check if in_dict)
    
    # Add summary paragraph
    summary = doc.add_paragraph()
    summary.add_run(f"Total words: {total_words}\n").bold = True
    summary.add_run(f"Root words found in dictionary: {found_words}\n").bold = True
    
    # List all words and their status
    for word, in_dict in dict_check:
        p = doc.add_paragraph(style='List Bullet')
        p.text = f"{word}: {'Found in dictionary' if in_dict else 'Not found'}"
    
    # Show final stemmed text
    doc.add_heading('5. Final Stemmed Text:', level=1)
    doc.add_paragraph(stemmed_text)
    
    # Add VSM calculations if available
    if vsm_data:
        doc.add_heading('6. Vector Space Model Calculations (TF Only):', level=1)
        
        # Document Term Matrix
        doc.add_heading('Term-Document Matrix (Term Frequencies):', level=2)
        terms = sorted(list(vsm_data['terms']))
        
        # Create table for term frequencies
        table = doc.add_table(rows=1, cols=len(vsm_data['documents']) + 1)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Term'
        for i in range(len(vsm_data['documents'])):
            header_cells[i+1].text = f'Doc {i+1}'
            
        # Add term frequencies
        for term in terms:
            row_cells = table.add_row().cells
            row_cells[0].text = term
            for doc_id in range(len(vsm_data['documents'])):
                freq = vsm_data['term_doc_freq'][term][doc_id]
                row_cells[doc_id+1].text = str(freq)
            
        # Document Vectors
        doc.add_heading('Term Frequency Vectors:', level=2)
        for i, doc_vector in enumerate(vsm_data['doc_vectors']):
            doc.add_paragraph(f'Document {i+1}:')
            vector_table = doc.add_table(rows=1, cols=2)
            vector_table.style = 'Table Grid'
            header_cells = vector_table.rows[0].cells
            header_cells[0].text = 'Term'
            header_cells[1].text = 'TF Weight'
            
            for term in terms:
                if term in doc_vector:
                    row_cells = vector_table.add_row().cells
                    row_cells[0].text = term
                    row_cells[1].text = str(doc_vector[term])

    doc.save(output_filename)
    return output_filename

class VSM:
    def __init__(self, stemmer):
        self.stemmer = stemmer
        self.documents = []
        self.doc_vectors = []
        self.terms = set()
        self.term_doc_freq = defaultdict(lambda: defaultdict(int))

    def add_document(self, doc_id, content):
        # Stem the document content
        stemmed_text, _ = self.stemmer.stem_text(content)
        tokens = stemmed_text.split()
        
        # Store document
        self.documents.append({
            'id': doc_id,
            'content': content,
            'stemmed': stemmed_text,
            'tokens': tokens
        })
        
        # Update term frequencies
        for term in tokens:
            self.terms.add(term)
            self.term_doc_freq[term][doc_id] += 1

    def calculate_weights(self):
        # Calculate TF vectors for each document (without IDF)
        self.doc_vectors = []
        for doc in self.documents:
            vector = {}
            doc_id = doc['id']
            for term in self.terms:
                tf = self.term_doc_freq[term][doc_id]
                if tf > 0:
                    # Using raw term frequency
                    vector[term] = tf
            self.doc_vectors.append(vector)

    def search(self, query):
        # Stem the query
        stemmed_query, _ = self.stemmer.stem_text(query)
        query_terms = stemmed_query.split()

        # Create query vector using TF only
        query_vector = {}
        for term in query_terms:
            if term in self.terms:
                query_vector[term] = query_terms.count(term)

        # Calculate similarities
        results = []
        for i, doc_vector in enumerate(self.doc_vectors):
            similarity = self.cosine_similarity(query_vector, doc_vector)
            results.append((self.documents[i], similarity))

        # Sort by similarity score
        results.sort(key=lambda x: x[1], reverse=True)
        return results

    def cosine_similarity(self, vec1, vec2):
        if not vec1 or not vec2:
            return 0
            
        dot_product = sum(vec1.get(term, 0) * vec2.get(term, 0) for term in self.terms)
        
        norm1 = math.sqrt(sum(value * value for value in vec1.values()))
        norm2 = math.sqrt(sum(value * value for value in vec2.values()))
        
        if norm1 == 0 or norm2 == 0:
            return 0
            
        return dot_product / (norm1 * norm2)

# Modify the main section to include VSM functionality
if __name__ == "__main__":
    stemmer = Stemmer()
    vsm = VSM(stemmer)
    
    root = tk.Tk()
    root.withdraw()
    
    # Allow multiple file selection
    file_paths = filedialog.askopenfilenames(
        title="Select Documents",
        filetypes=[
            ("All supported files", "*.txt;*.pdf;*.docx"),
            ("Text files", "*.txt"),
            ("PDF files", "*.pdf"),
            ("Word files", "*.docx"),
            ("All files", "*.*")
        ]
    )
    
    if file_paths:
        try:
            # Process all selected documents
            for i, file_path in enumerate(file_paths):
                text = read_file_content(file_path)
                vsm.add_document(i, text)
            
            # Calculate document vectors
            vsm.calculate_weights()
            
            # Prepare VSM data for export without IDF
            vsm_data = {
                'terms': vsm.terms,
                'documents': vsm.documents,
                'term_doc_freq': vsm.term_doc_freq,
                'doc_vectors': vsm.doc_vectors
            }
            
            # Process first document for stemming results
            text = read_file_content(file_paths[0])
            stemmed_text, steps = stemmer.stem_text(text)
            
            # Export results with VSM data
            output_file = export_to_word(text, stemmed_text, steps, file_paths[0], vsm_data)
            print(f"\nResults exported to: {output_file}")
            
            # Search loop
            while True:
                query = input("\nEnter search query (or 'quit' to exit): ")
                if query.lower() == 'quit':
                    break
                
                results = vsm.search(query)
                print("\nSearch Results:")
                for doc, score in results:
                    if score > 0:
                        print(f"\nDocument: {file_paths[doc['id']]}")
                        print(f"Similarity Score: {score:.4f}")
                        print("Preview:", doc['content'][:200], "...")
                
        except Exception as e:
            print(f"Error processing files: {e}")
    else:
        print("No files selected")
