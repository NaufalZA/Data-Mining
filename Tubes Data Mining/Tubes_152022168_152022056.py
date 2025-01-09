import csv
import re
from tkinter import filedialog
import tkinter as tk
import PyPDF2
from docx import Document
import os
import string
import math
from collections import defaultdict

class Stemmer:
    def __init__(self):
        self.kamus = self.load_kamus()
        self.stopwords = self.load_stopwords()
        self.forbidden_combinations = {
            'be': ['i'],
            'di': ['an'],
            'ke': ['i', 'kan'],
            'me': ['an'],
            'se': ['i', 'kan']
        }
        self.punctuation = (string.punctuation + '"' + '"' + ''' + ''' + '—' + '–' + 
                          '•' + '·' + '⋅' + '∙' + '‧' + '・' + '･' + '►' + '▪' + '○' + 
                          '●' + '♦' + '■' + '★' + '☆' + '✓' + '✔' + '❖')
        self.prefix_types = {
            'di': 'di-',
            'ke': 'ke-',
            'se': 'se-',
            'te': 'te-',
            'ter': 'ter-',
            'me': 'me-',
            'be': 'be-',
            'pe': 'pe-'
        }

    def load_kamus(self):
        with open('documents/Kamus.txt', 'r') as file:
            return set(word.strip().lower() for word in file)

    def load_stopwords(self):
        stopwords = set()
        with open('documents/Stopword.csv', 'r') as file:
            reader = csv.reader(file)
            for row in reader:
                stopwords.add(row[0].lower())
        return stopwords

    def is_valid_word(self, word):
        return bool(re.match('^[a-zA-Z]+$', word))

    def remove_stopwords(self, text):
        words = text.lower().split()
        return ' '.join([word for word in words if word not in self.stopwords])

    def check_kamus(self, word):
        return word.lower() in self.kamus

    def remove_inflection_suffixes(self, word):
        
        if word.endswith(('lah', 'kah', 'tah', 'pun')):
            word = re.sub(r'(lah|kah|tah|pun)$', '', word)
            
        
        if word.endswith(('ku', 'mu', 'nya')):
            word = re.sub(r'(ku|mu|nya)$', '', word)
            
        return word

    def remove_derivation_suffixes(self, word):
        if word.endswith(('i', 'an', 'kan')):
            original = word
            
            
            if word.endswith('kan'):
                word = word[:-3]
            elif word.endswith(('i', 'an')):
                word = word[:-2]

            
            if original.endswith('an') and word.endswith('k'):
                word = word[:-1]
                if self.check_kamus(word):
                    return word
                word = word + 'k'

            
            if self.check_kamus(word):
                return word
                
            
            return original
            
        return word

    def is_forbidden_combination(self, prefix, suffix):
        forbidden = {
            'be': ['i'],
            'di': ['an'],
            'ke': ['i', 'kan'],
            'me': ['an'],
            'se': ['i', 'kan']
        }
        return prefix in forbidden and suffix in forbidden[prefix]

    def remove_prefix(self, word, iteration=1):
        if iteration > 3:
            return word, None  

        
        previous_prefix = None if iteration == 1 else self.get_prefix_type(word)

        if word[:2] in ['di', 'ke', 'se']:
            prefix = word[:2]
            stemmed = word[2:]
            prefix_type = self.prefix_types[prefix]
            
        elif word.startswith('ter'):
            prefix = 'ter'
            stemmed = word[3:]
            prefix_type = 'ter-'
            
        elif word.startswith('te'):
            prefix = 'te'
            stemmed = word[2:]
            prefix_type = 'te-'
            
        elif word.startswith(('me', 'pe', 'be')):
            if len(word) > 3 and word[2] == 'r' and word[3] in 'aiueo':
                prefix = word[:3]
                stemmed = word[3:]
            else:
                prefix = word[:2]
                stemmed = word[2:]
            prefix_type = self.prefix_types[prefix[:2]]
        else:
            return word, None

        
        if previous_prefix == prefix_type:
            return word, None

        
        if self.check_kamus(stemmed):
            return stemmed, prefix_type

        
        recoded = self.recode_prefix(prefix, stemmed)
        if recoded != stemmed and self.check_kamus(recoded):
            return recoded, prefix_type
            
        
        next_word, next_prefix = self.remove_prefix(word, iteration + 1)
        return next_word, next_prefix or prefix_type

    def recode_prefix(self, prefix, word):
        """Handle special recoding cases"""
        if prefix in ['me', 'pe']:
            if word.startswith('ng'):
                return word[2:]  
            elif word.startswith('ny'):
                return 's' + word[2:]  
            elif word.startswith('n'):
                return 't' + word[1:]  
        return word

    def stem_word(self, word):
        steps = []
        if not self.is_valid_word(word):
            return word, steps
            
        
        steps.append(f"Step 1: Checking '{word}' in dictionary")
        if self.check_kamus(word):
            steps.append("Result: Found in dictionary")
            return word, steps

        original_word = word
        suffix_removed = False
        prefix_type = None

        
        steps.append("Step 2: Checking inflection suffixes")
        if any(word.endswith(suffix) for suffix in ['lah', 'kah', 'tah', 'pun', 'ku', 'mu', 'nya']):
            old_word = word
            word = self.remove_inflection_suffixes(word)
            steps.append(f"Removed inflection suffix: {old_word} → {word}")

        
        steps.append(f"Step 3: Checking derivation suffixes for '{word}'")
        if any(word.endswith(suffix) for suffix in ['i', 'an', 'kan']):
            temp_word = word
            if word.endswith('kan'):
                word = word[:-3]
                suffix_removed = 'kan'
                steps.append(f"Removed -kan: '{temp_word}' → '{word}'")
            elif word.endswith('i'):
                word = word[:-1]
                suffix_removed = 'i'
                steps.append(f"Removed -i: '{temp_word}' → '{word}'")
            elif word.endswith('an'):
                word = word[:-2]
                suffix_removed = 'an'
                steps.append(f"Removed -an: '{temp_word}' → '{word}'")
                
                if word.endswith('k'):
                    k_word = word[:-1]
                    steps.append(f"Checking k-removal: '{word}' → '{k_word}'")
                    if self.check_kamus(k_word):
                        steps.append(f"Result: Found '{k_word}' in dictionary")
                        return k_word, steps
                    word = temp_word
                    steps.append("Restored original: k-removal unsuccessful")

            if self.check_kamus(word):
                steps.append(f"Result: Found '{word}' in dictionary")
                return word, steps
            word = temp_word
            steps.append("Restored original: suffix removal unsuccessful")

        
        if suffix_removed:
            steps.append("Step 4a: Checking prefix-suffix combinations")
            prefix = self.get_prefix_type(word)
            if prefix and suffix_removed in self.forbidden_combinations.get(prefix, []):
                steps.append(f"Found forbidden combination: {prefix}- with -{suffix_removed}")
                return original_word, steps
        
        steps.append("Step 4b: Removing prefixes")
        word, prefix_type = self.remove_prefix(word)
        if prefix_type:
            steps.append(f"Removed prefix type: {prefix_type}")

        
        if prefix_type:
            steps.append("Step 5: Checking recoding rules")
            recoded = self.recode_prefix(prefix_type.rstrip('-'), word)
            if recoded != word:
                steps.append(f"Applied recoding: {word} → {recoded}")
                word = recoded

        
        steps.append("Step 6: No root word found, returning original word")
        return original_word, steps

    def get_prefix_type(self, word):
        if word.startswith(('di', 'ke', 'se')):
            return word[:2]
        elif word.startswith(('ter', 'bel')):
            return word[:3]
        elif word.startswith(('me', 'pe', 'be')):
            return word[:2]
        return None

    def tokenize(self, text):
        """
        Tokenize text into words while handling punctuation and special cases.
        Returns list of tokens and their positions.
        """
        
        text = text.lower()
        
        
        text = text.replace('\n', ' ')
        text = text.replace('\t', ' ')
        text = text.replace('\r', ' ')
        
        
        for p in self.punctuation:
            text = text.replace(p, ' ')
            
        
        text = re.sub(r'[\u200b\u200c\u200d\ufeff\xa0]', ' ', text)
        
        
        text = ' '.join(text.split())
        
        
        words = text.split()
        tokens = []
        position = 0
        
        for word in words:
            
            if (word and 
                not word.isnumeric() and 
                len(word) > 1 and 
                self.is_valid_word(word)):
                tokens.append({
                    'token': word,
                    'position': position,
                    'original': word
                })
            position += 1
            
        return tokens

    def stem_text(self, text):
        
        tokens = self.tokenize(text)
        
        
        tokens = [t for t in tokens if t['token'] not in self.stopwords]
        
        results = []
        all_steps = []
        
        
        for token in tokens:
            stemmed_word, steps = self.stem_word(token['token'])
            token['stemmed'] = stemmed_word
            results.append(stemmed_word)
            if token['token'] != stemmed_word:  
                all_steps.append((token['original'], stemmed_word, steps))
            
        return ' '.join(results), all_steps

def read_file_content(file_path):
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ''
            for page in pdf_reader.pages:
                text += page.extract_text() + '\n'
            return text
            
    elif file_extension == '.docx':
        doc = Document(file_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        return text
        
    elif file_extension == '.txt':
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    else:
        raise ValueError("Unsupported file format")

def export_to_word(original_text, stemmed_text, steps, input_file_path, vsm_data=None):
    original_filename = os.path.splitext(os.path.basename(input_file_path))[0]
    output_filename = f"results/Hasil_{original_filename}.docx"
    
    doc = Document()
    doc.add_heading('Hasil Stemming dan Perhitungan Temu Balik', 0)
    
    # Tokenisasi
    tokens = stemmer.tokenize(original_text)
    doc.add_heading('1. Hasil Tokenisasi:', level=1)
    token_text = ', '.join([t['token'] for t in tokens])
    doc.add_paragraph(token_text)
    
    # Angka
    valid_tokens = [t for t in tokens if not any(c.isdigit() for c in t['token'])]
    doc.add_heading('2. Hasil Setelah Removal Stopword dan Angka:', level=1)
    valid_text = ', '.join([t['token'] for t in valid_tokens])
    doc.add_paragraph(valid_text)
    
    # Kamus
    dict_check = [(t['token'], stemmer.check_kamus(t['token'])) for t in valid_tokens]
    doc.add_heading('3. Hasil Pengecekan Kata Dalam Kamus:', level=1)
    
    total_words = len(dict_check)
    found_words = sum(1 for _, in_dict in dict_check if in_dict)
    
    summary = doc.add_paragraph()
    summary.add_run(f"Total kata: {total_words}\n").bold = True
    summary.add_run(f"Kata dasar dalam kamus: {found_words}\n").bold = True
    
    for word, in_dict in dict_check:
        p = doc.add_paragraph(style='List Bullet')
        p.text = f"{word}: {'Ada dalam kamus' if in_dict else 'Tidak ada dalam kamus'}"
    
    doc.add_heading('4. Hasil Akhir Stemming:', level=1)
    doc.add_paragraph(stemmed_text)
    
    if vsm_data:
        doc.add_heading('5. Analisis VSM:', level=1)
        
        doc.add_paragraph("Berikut adalah matriks term-dokumen yang menunjukkan frekuensi kemunculan setiap kata pada setiap dokumen:")
        
        terms = sorted(list(vsm_data['terms']))
        table = doc.add_table(rows=1, cols=len(vsm_data['documents']) + 1)
        table.style = 'Table Grid'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Term'
        for i in range(len(vsm_data['documents'])):
            header_cells[i+1].text = f'Doc {i+1}'
            
        for term in terms:
            row_cells = table.add_row().cells
            row_cells[0].text = term
            for doc_id in range(len(vsm_data['documents'])):
                freq = vsm_data['term_doc_freq'][term][doc_id]
                row_cells[doc_id+1].text = str(freq)

        # Tambahkan tabel IDF
        doc.add_heading('Inverse Document Frequency (IDF):', level=2)
        idf_table = doc.add_table(rows=1, cols=2)
        idf_table.style = 'Table Grid'
        idf_table.rows[0].cells[0].text = 'Term'
        idf_table.rows[0].cells[1].text = 'IDF'
        
        for term in sorted(vsm_data['terms']):
            row_cells = idf_table.add_row().cells
            row_cells[0].text = term
            row_cells[1].text = f"{vsm_data['idf'][term]:.4f}"

        # Tambahkan tabel TF-IDF
        doc.add_heading('TF-IDF Weights:', level=2)
        tf_idf_table = doc.add_table(rows=1, cols=len(vsm_data['documents']) + 1)
        tf_idf_table.style = 'Table Grid'
        
        header_cells = tf_idf_table.rows[0].cells
        header_cells[0].text = 'Term'
        for i in range(len(vsm_data['documents'])):
            header_cells[i+1].text = f'Doc {i+1}'
            
        for term in sorted(list(vsm_data['terms'])):
            row_cells = tf_idf_table.add_row().cells
            row_cells[0].text = term
            for doc_id in range(len(vsm_data['documents'])):
                tf_idf = vsm_data['tf_idf_vectors'][doc_id].get(term, 0)
                row_cells[doc_id+1].text = f"{tf_idf:.4f}"

    if vsm_data and 'query' in vsm_data:
        doc.add_heading('6. Hasil Pencarian dan Perhitungan Similiarity:', level=1)
        doc.add_paragraph(f"Kata Kunci: {vsm_data['query']}")
        
        doc.add_heading('Analisis Keyword:', level=2)
        doc.add_paragraph("Frekuensi kemunculan kata pada query setelah preprocessing:")
        
        stemmed_query, _ = stemmer.stem_text(vsm_data['query'])
        query_terms = stemmed_query.split()
        query_vector = {}
        for term in query_terms:
            if term in vsm_data['terms']:
                query_vector[term] = query_terms.count(term)
                
        query_table = doc.add_table(rows=1, cols=2)
        query_table.style = 'Table Grid'
        query_table.rows[0].cells[0].text = 'Term'
        query_table.rows[0].cells[1].text = 'Frequency'
        for term, freq in query_vector.items():
            row = query_table.add_row()
            row.cells[0].text = term
            row.cells[1].text = str(freq)
            
        doc.add_heading('Perhitungan Kemiripan:', level=2)
        for i, (result_doc, score) in enumerate(vsm_data['search_results']):
            if score > 0:
                doc_id = result_doc['id'] + 1
                doc.add_paragraph(f"\nSimilarity (D{doc_id}, Q) = cos(D{doc_id}, Q)")
                
                doc_vector = vsm_data['doc_vectors'][result_doc['id']]
                query_vector = {term: query_terms.count(term) for term in query_terms if term in vsm_data['terms']}
                
                doc.add_paragraph("Perhitungan Dot:")
                dot_products = []
                dot_product = 0
                for term in vsm_data['terms']:
                    q_val = query_vector.get(term, 0)
                    d_val = doc_vector.get(term, 0)
                    if q_val > 0 or d_val > 0:
                        dot_products.append(f"({q_val} × {d_val})")
                        dot_product += q_val * d_val
                
                doc.add_paragraph("Dot(D{}, Q) = {}".format(
                    doc_id,
                    " + ".join(dot_products) + f" = {dot_product}"
                ))
                
                doc.add_paragraph("\nPerhitungan Magnitude :")
                q_magnitude_parts = []
                for term, val in query_vector.items():
                    if val > 0:
                        q_magnitude_parts.append(f"{val}²")
                query_magnitude = math.sqrt(sum(v * v for v in query_vector.values()))
                doc.add_paragraph("|Q| = √({}) = {:.4f}".format(
                    " + ".join(q_magnitude_parts),
                    query_magnitude
                ))
                
                d_magnitude_parts = []
                for term, val in doc_vector.items():
                    if val > 0:
                        d_magnitude_parts.append(f"{val}²")
                doc_magnitude = math.sqrt(sum(v * v for v in doc_vector.values()))
                doc.add_paragraph("|D{}| = √({}) = {:.4f}".format(
                    doc_id,
                    " + ".join(d_magnitude_parts),
                    doc_magnitude
                ))
                
                # doc.add_paragraph(f"\nFinal Similarity Calculation:")
                doc.add_paragraph("\nSimilarity(D{}, Q) = Dot(D{}, Q) / (|D{}| × |Q|)".format(
                    doc_id, doc_id, doc_id
                ))
                doc.add_paragraph("Similarity(D{}, Q) = {} / ({:.4f} × {:.4f})".format(
                    doc_id, dot_product, doc_magnitude, query_magnitude
                ))
                doc.add_paragraph(f"Similarity(D{doc_id}, Q) = {score:.4f}")
                
                doc.add_paragraph("-" * 40)

        doc.add_heading('Ranking Documents:', level=2)
        results_table = doc.add_table(rows=1, cols=2)
        results_table.style = 'Table Grid'
        header_cells = results_table.rows[0].cells
        header_cells[0].text = 'Judul Dokumen'
        header_cells[1].text = 'Nilai Similiarity'
        
        for result_doc, score in vsm_data['search_results']:
            if score > 0:
                row_cells = results_table.add_row().cells
                row_cells[0].text = os.path.basename(file_paths[result_doc['id']])
                row_cells[1].text = f"{score:.4f}"

    doc.save(output_filename)
    return output_filename

class VSM:
    def __init__(self, stemmer):
        self.stemmer = stemmer
        self.documents = []
        self.doc_vectors = []
        self.terms = set()
        self.term_doc_freq = defaultdict(lambda: defaultdict(int))
        self.idf = {}  # Menambahkan variabel untuk menyimpan nilai IDF
        self.tf_idf_vectors = []  # Menambahkan variabel untuk menyimpan vektor TF-IDF

    def add_document(self, doc_id, content):
        
        stemmed_text, _ = self.stemmer.stem_text(content)
        tokens = stemmed_text.split()
        
        
        self.documents.append({
            'id': doc_id,
            'content': content,
            'stemmed': stemmed_text,
            'tokens': tokens
        })
        
        
        for term in tokens:
            self.terms.add(term)
            self.term_doc_freq[term][doc_id] += 1

    def calculate_idf(self):
        """Menghitung IDF untuk setiap term"""
        N = len(self.documents)
        for term in self.terms:
            df = sum(1 for doc_id in range(N) if self.term_doc_freq[term][doc_id] > 0)
            # Add 1 to prevent log(1) = 0, this is called "smooth IDF"
            self.idf[term] = math.log10((N + 1) / (df + 1)) + 1

    def calculate_weights(self):
        """Menghitung bobot TF-IDF"""
        self.calculate_idf()
        self.doc_vectors = []
        self.tf_idf_vectors = []
        
        for doc in self.documents:
            vector = {}
            tf_idf_vector = {}
            doc_id = doc['id']
            
            for term in self.terms:
                tf = self.term_doc_freq[term][doc_id]
                if tf > 0:
                    # Use raw TF instead of log TF for better results with small documents
                    weighted_tf = tf
                    vector[term] = tf
                    tf_idf_vector[term] = weighted_tf * self.idf[term]
                    
            self.doc_vectors.append(vector)
            self.tf_idf_vectors.append(tf_idf_vector)

    def search(self, query):
        
        stemmed_query, _ = self.stemmer.stem_text(query)
        query_terms = stemmed_query.split()

        # Hitung TF-IDF untuk query
        query_vector = {}
        query_tf_idf = {}
        for term in query_terms:
            if term in self.terms:
                tf = query_terms.count(term)
                query_vector[term] = tf
                # Use raw TF for query as well
                query_tf_idf[term] = tf * self.idf.get(term, 0)

        
        results = []
        for i, doc_vector in enumerate(self.tf_idf_vectors):
            similarity = self.cosine_similarity(query_tf_idf, doc_vector)
            results.append((self.documents[i], similarity))

        
        results.sort(key=lambda x: x[1], reverse=True)
        return results

    def cosine_similarity(self, vec1, vec2):
        if not vec1 or not vec2:
            return 0
            
        dot_product = sum(vec1.get(term, 0) * vec2.get(term, 0) for term in self.terms)
        
        norm1 = math.sqrt(sum(value * value for value in vec1.values()))
        norm2 = math.sqrt(sum(value * value for value in vec2.values()))
        
        if norm1 == 0 or norm2 == 0:
            return 0
            
        return dot_product / (norm1 * norm2)


if __name__ == "__main__":
    stemmer = Stemmer()
    vsm = VSM(stemmer)
    
    root = tk.Tk()
    root.withdraw()
    
    file_paths = filedialog.askopenfilenames(
        title="Select Documents",
        filetypes=[
            ("All supported files", "*.txt;*.pdf;*.docx"),
            ("Text files", "*.txt"),
            ("PDF files", "*.pdf"),
            ("Word files", "*.docx"),
            ("All files", "*.*")
        ]
    )
    
    if file_paths:
        try:
            # Load documents
            for i, file_path in enumerate(file_paths):
                text = read_file_content(file_path)
                vsm.add_document(i, text)
                print(f"Dokumen {i+1}: {os.path.basename(file_path)}")
            
            # Calculate weights
            vsm.calculate_weights()
            
            while True:
                query = input("\nMasukkan kata kunci pencarian (ketik 'keluar' untuk berhenti): ")
                if query.lower() == 'keluar':
                    break
                
                results = vsm.search(query)
                print("\nHasil Pencarian:")
                print("-" * 50)
                
                # Export results for all documents
                for doc_id in range(len(file_paths)):
                    doc_path = file_paths[doc_id]
                    score = next((score for doc, score in results if doc['id'] == doc_id), 0)
                    print(f"Dokumen {doc_id + 1}: {os.path.basename(doc_path)}")
                    print(f"Nilai Kemiripan: {score:.4f}\n")
                    
                    text = read_file_content(doc_path)
                    stemmed_text, steps = stemmer.stem_text(text)
                    vsm_data = {
                        'terms': vsm.terms,
                        'documents': vsm.documents,
                        'term_doc_freq': vsm.term_doc_freq,
                        'doc_vectors': vsm.doc_vectors,
                        'tf_idf_vectors': vsm.tf_idf_vectors,
                        'idf': vsm.idf,
                        'query': query,
                        'search_results': results
                    }
                    export_to_word(text, stemmed_text, steps, doc_path, vsm_data)
                print("-" * 50)
                
        except Exception as e:
            print(f"Terjadi kesalahan: {e}")
    else:
        print("Tidak ada file yang dipilih")
